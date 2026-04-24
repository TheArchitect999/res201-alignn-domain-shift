#!/usr/bin/env python3
"""
RES201 Final Document Builder
Builds three Word documents from polished markdown manuscripts using python-docx.
Inserts citations, figures, and inline formatting. Matches JURI template styles.
"""

import json, csv, re, os
from collections import OrderedDict
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# ==============================================================================
# CONFIGURATION
# ==============================================================================
BASE = r'c:\Users\lenovo\res201-alignn-domain-shift'
FIGS_DIR = os.path.join(BASE, r'reports\final_paper_factory\02_figure_memos\core_figures')
OUTPUT = os.path.join(BASE, 'reports', 'final_paper_factory', '05_final_docx')
os.makedirs(OUTPUT, exist_ok=True)

TEMPLATE = os.path.join(BASE, 'paper_sources', 'JURI_Template.docx')
REFS_JSON = os.path.join(OUTPUT, 'res201_final_refs.json')
CITATION_MAP = os.path.join(BASE, 'reports', 'final_paper_factory', '04_drafts',
                             'phase12_full_manuscripts', 'res201_citation_map_final.csv')

MANUSCRIPT_FILES = {
    'oxide':    os.path.join(BASE, 'reports', 'final_paper_factory', '04_drafts',
                             'phase12_full_manuscripts', 'oxide_polished_v3.md'),
    'nitride':  os.path.join(BASE, 'reports', 'final_paper_factory', '04_drafts',
                             'phase12_full_manuscripts', 'nitride_polished_v3.md'),
    'combined': os.path.join(BASE, 'reports', 'final_paper_factory', '04_drafts',
                             'phase12_full_manuscripts', 'combined_paper_polished_v4.md'),
}
MANUSCRIPT_KEYS = {
    'oxide': 'oxide_report',
    'nitride': 'nitride_report',
    'combined': 'combined_report',
}
OUTPUT_FILES = {
    'oxide':    os.path.join(OUTPUT, 'oxide_final.docx'),
    'nitride':  os.path.join(OUTPUT, 'nitride_final.docx'),
    'combined': os.path.join(OUTPUT, 'combined_final.docx'),
}
AUTHORS = {
    'oxide': {
        'line': 'Muhammad Ali Bin Sarwar',
        'orcid': '0009-0009-8398-0430',
        'email': 's202432700@kfupm.edu.sa',
        'affiliation': '[Affiliation — confirm with authors; email domain: kfupm.edu.sa]',
    },
    'nitride': {
        'line': 'Faizan Ahmed',
        'orcid': '0009-0003-1857-9348',
        'email': 's202457920@kfupm.edu.sa',
        'affiliation': '[Affiliation — confirm with authors; email domain: kfupm.edu.sa]',
    },
    'combined': {
        'line': 'Faizan Ahmed, Muhammad Ali Bin Sarwar',
        'orcid': '0009-0003-1857-9348 (F.A.); 0009-0009-8398-0430 (M.A.B.S.)',
        'email': 's202457920@kfupm.edu.sa',
        'affiliation': '[Affiliation — confirm with authors; email domain: kfupm.edu.sa]',
    },
}
ACK_TEXT = ('The authors gratefully acknowledge use of the JARVIS infrastructure '
            '(https://jarvis.nist.gov) for datasets, pretrained models, and benchmark splits '
            'used throughout this work.')

# Figure file map: manuscript marker ID -> filename in core_figures
FIGURE_FILES = {
    'FIG_SCHEMATIC':              'FIG_SCHEMATIC.png',
    'FIG_ZS_COMPARISON':          'FIG_ZS_COMPARISON.png',
    'FIG_TRANSFER_BENEFIT':       'FIG_TRANSFER_BENEFIT.png',
    'FIG_S1_LC_OXIDE':            'FIG_S1_LC_OXIDE.png',
    'FIG_S1_LC_NITRIDE':          'FIG_S1_LC_NITRIDE.png',
    'FIG_S1_COMP_OXIDE':          'FIG_S1_COMP_OXIDE.png',
    'FIG_S1_COMP_NITRIDE':        'FIG_S1_COMP_NITRIDE.png',
    'FIG_S1_PARITY_OXIDE_N10':    'FIG_S1_PARITY_OXIDE_N10.png',
    'FIG_S1_PARITY_OXIDE_N1000':  'FIG_S1_PARITY_OXIDE_N1000.png',
    'FIG_S1_PARITY_NITRIDE_N10':  'FIG_S1_PARITY_NITRIDE_N10.png',
    'FIG_S1_PARITY_NITRIDE_N1000':'FIG_S1_PARITY_NITRIDE_N1000.png',
    'FIG_EA_6A_PCA':              'FIG_EA_6A_PCA.png',
    'FIG_EA_6B_TSNE':             'FIG_EA_6B_TSNE_P30.png',
    'FIG_EA_6C_UMAP':             'FIG_EA_6C_UMAP_N30.png',
    'FIG_EA_6D_BOXPLOT':          'FIG_EA_6D_KNN5_BOXPLOT.png',
    'FIG_EA_6D_SCATTER':          'FIG_EA_6D_KNN5_SCATTER.png',
}
FIG_CAPTIONS = {
    'FIG_SCHEMATIC':
        'Study design overview. The pretrained formation-energy ALIGNN model is evaluated '
        'zero-shot on both oxide (in-distribution) and nitride (out-of-distribution) test sets, '
        'then fine-tuned across labelled-data sizes N ∈ {10, 50, 100, 200, 500, 1000} with five '
        'random seeds each, and compared to matched from-scratch baselines at N = 50 and N = 500.',
    'FIG_ZS_COMPARISON':
        'Zero-shot MAE comparison. Same pretrained checkpoint evaluated on the 1,484-structure '
        'oxide test set (0.0342 eV/atom) and the 242-structure nitride test set (0.0695 eV/atom) '
        'without any target-family fine-tuning. The nitride error is 2.03x the oxide error.',
    'FIG_TRANSFER_BENEFIT':
        'Transfer benefit across families. From-scratch minus fine-tune mean test MAE at N = 50 '
        'and N = 500 for oxides (blue) and nitrides (orange). Note: the nitride N = 50 bar '
        'reflects pretrained-initialization advantage, not fine-tuning adaptation '
        '(mean_best_epoch = 1.0). Error bars show ±1 SD across five seeds.',
    'FIG_S1_LC_OXIDE':
        'Oxide Set 1 fine-tuning learning curve. Mean test MAE ± 1 SD (five seeds) at each '
        'labelled-data size. Dashed line: pretrained zero-shot MAE (0.0342 eV/atom). '
        'Fine-tuning converges toward but does not cross the zero-shot benchmark.',
    'FIG_S1_LC_NITRIDE':
        'Nitride Set 1 fine-tuning learning curve. Mean test MAE ± 1 SD (five seeds). '
        'Dashed line: nitride zero-shot MAE (0.0695 eV/atom). N ≤ 200 rows are operationally '
        'inert (mean_best_epoch = 1.0 at all seeds); genuine adaptation begins at N = 500.',
    'FIG_S1_COMP_OXIDE':
        'Oxide Set 1 pretrained fine-tuning vs from-scratch at N = 50 and N = 500. '
        'Zero-shot reference shown. Fine-tune MAE is 6–10x lower than from-scratch MAE.',
    'FIG_S1_COMP_NITRIDE':
        'Nitride Set 1 pretrained fine-tuning vs from-scratch at N = 50 and N = 500. '
        'Zero-shot reference shown. The N = 50 difference reflects initialization advantage; '
        'the N = 500 difference is a clean adapted-vs-scratch comparison.',
    'FIG_S1_PARITY_OXIDE_N10':
        'Oxide parity plot at N = 10. Seed-averaged predictions vs DFT formation energy '
        '(1,484 test structures). On-figure MAE 0.0391 eV/atom, R² 0.9944. '
        'Effective zero-shot snapshot (mean_best_epoch = 1.0).',
    'FIG_S1_PARITY_OXIDE_N1000':
        'Oxide parity plot at N = 1,000. Seed-averaged predictions vs DFT formation energy '
        '(1,484 test structures). On-figure MAE 0.0383 eV/atom, R² 0.9943. '
        'Genuinely optimized run; per-seed SD narrows to 0.0053 eV/atom.',
    'FIG_S1_PARITY_NITRIDE_N10':
        'Nitride parity plot at N = 10. Seed-averaged predictions vs DFT formation energy '
        '(242 test structures). On-figure MAE 0.0828 eV/atom, R² 0.9841. '
        'Effective zero-shot snapshot (mean_best_epoch = 1.0).',
    'FIG_S1_PARITY_NITRIDE_N1000':
        'Nitride parity plot at N = 1,000. Seed-averaged predictions vs DFT formation energy '
        '(242 test structures). On-figure MAE 0.0829 eV/atom, R² 0.9837. '
        'Best-adapted configuration; still above the nitride zero-shot baseline (0.0695 eV/atom).',
    'FIG_EA_6A_PCA':
        'PCA projection of frozen last_alignn_pool embeddings (balanced pool, 4,092 structures). '
        'PC1 18.13%, PC2 9.47% explained variance. Oxide and nitride points occupy distinct '
        'regions. Descriptive visual support only; all inferential claims are from raw 256-D space.',
    'FIG_EA_6B_TSNE':
        't-SNE projection of frozen last_alignn_pool embeddings (balanced pool; perplexity = 30; '
        'standardized inputs). Oxide and nitride clusters are visually distinct. t-SNE preserves '
        'local neighbourhoods; global inter-cluster distances are not metric quantities.',
    'FIG_EA_6C_UMAP':
        'UMAP projection of frozen last_alignn_pool embeddings (balanced pool; n_neighbors = 30, '
        'min_dist = 0.1; standardized inputs). Consistent with raw-space separation metrics. '
        'Descriptive visual support only.',
    'FIG_EA_6D_BOXPLOT':
        'Distance–error boxplot. Hard (top 20% absolute zero-shot error, n = 49) vs easy '
        '(bottom 20%, n = 49) nitride test structures by 5-nearest-oxide-reference Euclidean '
        'distance in raw 256-D last_alignn_pool space. Hard-minus-easy mean gap: 0.8168 '
        '(95% CI 0.4746–1.1597; FDR q = 1.8 × 10⁻⁴).',
    'FIG_EA_6D_SCATTER':
        'Distance–error scatter plot. 5-nearest-oxide-reference distance vs absolute zero-shot '
        'error across all 242 nitride test structures in raw 256-D last_alignn_pool space. '
        'Spearman ρ = 0.3428 (95% CI 0.2214–0.4597; FDR q = 1.3 × 10⁻⁴). '
        'Association is correlational; see text for interpretation caveats.',
}

# ==============================================================================
# REFERENCE FORMATTING (Nature style)
# ==============================================================================
def fmt_initial(given):
    parts = given.split()
    return ' '.join(p[0] + '.' for p in parts if p)

def fmt_author(a):
    family = a.get('family', '')
    given = a.get('given', '')
    particle = a.get('non-dropping-particle', '')
    initials = fmt_initial(given)
    if particle:
        return f'{particle} {family}, {initials}'
    return f'{family}, {initials}'

def fmt_authors(authors):
    if not authors:
        return ''
    fa = [fmt_author(a) for a in authors]
    if len(fa) == 1:
        return fa[0]
    if len(fa) == 2:
        return f'{fa[0]} & {fa[1]}'
    if len(fa) <= 6:
        return ', '.join(fa[:-1]) + f' & {fa[-1]}'
    return f'{fa[0]} et al.'

def fmt_ref_structured(item):
    """Return list of (text, italic) segments for a Nature-style reference."""
    authors_str = fmt_authors(item.get('author', []))
    title = item.get('title', '')
    journal = item.get('journalAbbreviation', item.get('container-title', ''))
    volume = item.get('volume', '')
    pages = item.get('page', '')
    issued = item.get('issued', {}).get('date-parts', [['']])[0]
    year = str(issued[0]) if issued else ''
    note = item.get('note', '')
    item_type = item.get('type', 'article-journal')

    segs = []  # (text, italic)
    if item_type == 'article':  # preprint / arXiv
        m = re.search(r'arXiv:(\d{4}\.\d+)', note)
        arxiv_id = m.group(1) if m else ''
        url = f'https://arxiv.org/abs/{arxiv_id}' if arxiv_id else item.get('URL', '')
        segs.append((f'{authors_str}. {title}. Preprint at {url} ({year}).', False))
    else:
        segs.append((f'{authors_str}. {title}. ', False))
        if journal:
            segs.append((journal, True))  # journal in italics
            rest = ''
            if volume:
                rest += f' {volume}'
            if pages:
                rest += f', {pages}'
            if year:
                rest += f' ({year}).'
            segs.append((rest, False))
        elif year:
            segs.append((f'({year}).', False))
    return segs

# ==============================================================================
# CITATION MAP & TRACKER
# ==============================================================================
def load_citation_map(csv_path, ms_key):
    """Returns dict: placeholder_text -> [zotero_ids]"""
    result = {}
    with open(csv_path, newline='', encoding='utf-8') as f:
        for row in csv.DictReader(f):
            if row['manuscript'] != ms_key:
                continue
            ph = row['placeholder_text']
            ids = [z.strip() for z in row['zotero_key_or_id'].split('|')]
            if ph not in result:
                result[ph] = ids
    return result

class CitationTracker:
    def __init__(self, zotero_by_id, cite_map):
        self.zby = zotero_by_id
        self.cmap = cite_map
        self.num_by_id = OrderedDict()
        self._next = 1

    def _assign(self, ids):
        nums = []
        for id_ in ids:
            if id_ not in self.num_by_id:
                self.num_by_id[id_] = self._next
                self._next += 1
            nums.append(self.num_by_id[id_])
        return nums

    def replace(self, text):
        def repl(m):
            ph = m.group(0)
            ids = self.cmap.get(ph, [])
            if not ids:
                return '[?]'
            nums = self._assign(ids)
            return '[' + ','.join(str(n) for n in nums) + ']'
        return re.sub(r'\[CITE:[^\]]+\]', repl, text)

    def ref_list(self):
        """Returns list of (num, structured_segs) in citation order."""
        result = []
        for id_, num in self.num_by_id.items():
            item = self.zby.get(id_)
            if item:
                result.append((num, fmt_ref_structured(item)))
            else:
                result.append((num, [(f'Reference {id_} not in library.', False)]))
        return sorted(result)

# ==============================================================================
# INLINE FORMATTING PARSER
# ==============================================================================
def parse_inline(text):
    """Parse **bold**, *italic*, `code` in text. Returns list of (str, bold, italic, code)."""
    segs = []
    i = 0
    cur = ''
    state = {'bold': False, 'italic': False, 'code': False}

    def flush(s=state):
        nonlocal cur
        if cur:
            segs.append((cur, s['bold'], s['italic'], s['code']))
            cur = ''

    while i < len(text):
        c = text[i]
        if text[i:i+2] == '**' and not state['code']:
            flush()
            state['bold'] = not state['bold']
            i += 2
        elif c == '*' and text[i:i+2] != '**' and not state['code']:
            flush()
            state['italic'] = not state['italic']
            i += 1
        elif c == '`':
            flush()
            state['code'] = not state['code']
            i += 1
        else:
            cur += c
            i += 1
    flush()
    return segs

def add_inline(para, text):
    """Add text with inline markdown formatting to a Word paragraph."""
    for seg, bold, italic, code in parse_inline(text):
        if not seg:
            continue
        run = para.add_run(seg)
        run.bold = bold
        run.italic = italic
        if code:
            run.font.name = 'Courier New'
            run.font.size = Pt(9)

# ==============================================================================
# TABLE PARSER
# ==============================================================================
def parse_md_table(lines):
    """Parse markdown table lines into list-of-lists (skipping separator rows)."""
    rows = []
    for ln in lines:
        if re.match(r'^\s*\|[-:\s|]+\|?\s*$', ln):
            continue  # separator row
        cells = [c.strip() for c in re.split(r'(?<!\\)\|', ln.strip().strip('|'))]
        if cells:
            rows.append(cells)
    return rows

def add_md_table(doc, rows):
    """Add a Word table from parsed row data."""
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=ncols)
    tbl.style = 'Table Grid'
    for ri, row in enumerate(rows):
        for ci, cell_text in enumerate(row):
            if ci >= ncols:
                break
            cell = tbl.cell(ri, ci)
            cell.text = ''
            p = cell.paragraphs[0]
            p.clear()
            add_inline(p, cell_text)
            if ri == 0:
                for run in p.runs:
                    run.bold = True
    doc.add_paragraph()

# ==============================================================================
# FIGURE INSERTER
# ==============================================================================
def add_figure(doc, fig_id, fig_num, log):
    fname = FIGURE_FILES.get(fig_id)
    path = os.path.join(FIGS_DIR, fname) if fname else None
    if path and os.path.exists(path):
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        try:
            run.add_picture(path, width=Inches(5.8))
            log.append(f'  FIG {fig_num} ({fig_id}): inserted from {fname}')
        except Exception as e:
            log.append(f'  FIG {fig_num} ({fig_id}): picture error — {e}')
            para.text = f'[Figure {fig_num} — {fig_id} — insert manually]'
    else:
        p = doc.add_paragraph(f'[Figure {fig_num} — {fig_id} — image file not found]')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        log.append(f'  FIG {fig_num} ({fig_id}): MISSING — path not found')
    cap_text = FIG_CAPTIONS.get(fig_id, f'{fig_id}.')
    cap = doc.add_paragraph(f'Figure {fig_num}. {cap_text}')
    cap.style = doc.styles['Caption']
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

# ==============================================================================
# CORE DOCUMENT BUILDER
# ==============================================================================
def build_doc(ms_name, tracker, author_info, log):
    # Read manuscript
    with open(MANUSCRIPT_FILES[ms_name], encoding='utf-8') as f:
        raw = f.read()

    # Start from template (preserves styles + page layout)
    doc = Document(TEMPLATE)
    body = doc.element.body
    sect_pr = body.find(qn('w:sectPr'))
    # Remove all body elements except sectPr
    for el in list(body):
        if el is not sect_pr:
            body.remove(el)

    lines = raw.split('\n')
    in_table = False
    table_lines = []
    fig_num = 0
    state = 'pre_title'   # pre_title / pre_author / pre_affil / pre_email / body

    def flush_table():
        nonlocal in_table, table_lines
        rows = parse_md_table(table_lines)
        add_md_table(doc, rows)
        in_table = False
        table_lines = []

    i = 0
    while i < len(lines):
        line = lines[i]

        # --- Collect markdown table rows ---
        if re.match(r'^\s*\|', line):
            in_table = True
            table_lines.append(line)
            i += 1
            continue
        elif in_table:
            flush_table()
            table_lines = []

        stripped = line.strip()

        # --- Empty line ---
        if not stripped:
            i += 1
            continue

        # --- Horizontal rule ---
        if stripped == '---':
            i += 1
            continue

        # --- Figure marker ---
        fm = re.match(r'\[INSERT FIGURE (FIG_[\w]+) HERE\]', stripped)
        if fm:
            fig_id = fm.group(1)
            fig_num += 1
            add_figure(doc, fig_id, fig_num, log)
            i += 1
            continue

        # --- Table placeholder ---
        tm = re.match(r'\[INSERT TABLE (TAB_[\w]+) HERE\]', stripped)
        if tm:
            tab_id = tm.group(1)
            p = doc.add_paragraph(f'[TABLE: {tab_id} — formatted data table to be inserted here]')
            p.style = doc.styles['Caption']
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()
            log.append(f'  TABLE placeholder: {tab_id}')
            i += 1
            continue

        # --- Skip acknowledgements & references placeholder text ---
        if ('[ACKNOWLEDGEMENTS PLACEHOLDER' in stripped or
                '[REFERENCES PLACEHOLDER' in stripped):
            i += 1
            continue

        # --- Title (# heading) ---
        if re.match(r'^# (?!#)', line):
            title_text = line[2:].strip()
            p = doc.add_paragraph()
            p.style = doc.styles['Title']
            add_inline(p, title_text)
            state = 'pre_author'
            i += 1
            continue

        # --- Author line replacement ---
        if state == 'pre_author' and ('First Author Name' in stripped or
                                       'Second Author Name' in stripped):
            p = doc.add_paragraph()
            p.style = doc.styles['Normal']
            r = p.add_run(author_info['line'])
            state = 'pre_affil'
            i += 1
            continue

        if state == 'pre_affil' and '[Affiliation placeholder]' in stripped:
            p = doc.add_paragraph(author_info['affiliation'])
            p.style = doc.styles['Normal']
            state = 'pre_email'
            i += 1
            continue

        if state == 'pre_email' and '[Corresponding author email placeholder]' in stripped:
            p = doc.add_paragraph()
            p.style = doc.styles['Normal']
            p.add_run(f'¹Corresponding author: {author_info["email"]}').italic = True
            # ORCID on separate line
            doc.add_paragraph(f'ORCID: {author_info["orcid"]}').style = doc.styles['Normal']
            state = 'body'
            i += 1
            continue

        # --- Headings ---
        if line.startswith('#### '):
            txt = line[5:].strip()
            doc.add_heading(tracker.replace(txt), level=3)
            i += 1
            continue
        if line.startswith('### '):
            txt = line[4:].strip()
            doc.add_heading(tracker.replace(txt), level=2)
            i += 1
            continue
        if line.startswith('## '):
            txt = line[3:].strip()
            # Handle Acknowledgements heading
            if txt.lower().startswith('acknowledgement'):
                doc.add_heading('Acknowledgements', level=1)
                # Consume until we hit next heading or EOF, inserting ack text
                i += 1
                # Skip placeholder lines, add real text
                ack_added = False
                while i < len(lines):
                    nxt = lines[i].strip()
                    if nxt.startswith('## ') or nxt.startswith('# '):
                        break
                    if '[ACKNOWLEDGEMENTS PLACEHOLDER' in nxt or not nxt:
                        i += 1
                        continue
                    i += 1
                p = doc.add_paragraph(ACK_TEXT)
                p.style = doc.styles['Normal']
                continue
            # Handle References heading
            if txt.lower().startswith('reference'):
                doc.add_heading('References', level=1)
                i += 1
                # Skip placeholder, then insert reference list
                while i < len(lines):
                    nxt = lines[i].strip()
                    if nxt.startswith('#') or (nxt and not nxt.startswith('[REFERENCES')):
                        break
                    i += 1
                # Insert references
                for num, segs in tracker.ref_list():
                    p = doc.add_paragraph(style='Normal')
                    p.add_run(f'{num}. ')  # number + en-space
                    for seg_text, italic in segs:
                        run = p.add_run(seg_text)
                        run.italic = italic
                log.append(f'  References: inserted {len(tracker.ref_list())} entries')
                continue
            # Standard heading
            doc.add_heading(tracker.replace(txt), level=1)
            i += 1
            continue

        # --- Numbered / bulleted list items ---
        li_m = re.match(r'^(\d+\.)\s+(.*)', line)
        bl_m = re.match(r'^[-*]\s+(.*)', line)
        if li_m:
            p = doc.add_paragraph(style='List Paragraph')
            add_inline(p, tracker.replace(li_m.group(1) + ' ' + li_m.group(2)))
            i += 1
            continue
        if bl_m:
            p = doc.add_paragraph(style='List Paragraph')
            p.add_run('• ')
            add_inline(p, tracker.replace(bl_m.group(1)))
            i += 1
            continue

        # --- Normal paragraph ---
        processed = tracker.replace(stripped)
        p = doc.add_paragraph(style='Normal')
        add_inline(p, processed)
        i += 1

    # Flush any trailing table
    if in_table and table_lines:
        flush_table()

    return doc

# ==============================================================================
# MAIN
# ==============================================================================
def main():
    # Load Zotero library
    with open(REFS_JSON, encoding='utf-8') as f:
        refs = json.load(f)
    zotero_by_id = {}
    for item in refs:
        # Extract short key from URL id (last path segment)
        short_id = item['id'].rstrip('/').split('/')[-1]
        zotero_by_id[short_id] = item

    build_log = {}

    for ms_name in ('oxide', 'nitride', 'combined'):
        print(f'\n=== Building {ms_name} ===')
        log = []

        # Load citation map for this manuscript
        cite_map = load_citation_map(CITATION_MAP, MANUSCRIPT_KEYS[ms_name])
        log.append(f'Citation map: {len(cite_map)} unique placeholders')

        # Build tracker
        tracker = CitationTracker(zotero_by_id, cite_map)

        # Build document
        doc = build_doc(ms_name, tracker, AUTHORS[ms_name], log)

        # Save
        out_path = OUTPUT_FILES[ms_name]
        doc.save(out_path)
        n_refs = len(tracker.ref_list())
        log.append(f'Saved: {out_path}')
        log.append(f'Total citations used: {n_refs}')
        print(f'  Saved {out_path}  ({n_refs} references)')

        build_log[ms_name] = log

    # Print summary
    print('\n=== Build complete ===')
    for ms, entries in build_log.items():
        print(f'\n{ms.upper()}:')
        for e in entries:
            print(f'  {e}')

    return build_log

if __name__ == '__main__':
    main()
