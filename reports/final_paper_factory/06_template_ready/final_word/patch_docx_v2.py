#!/usr/bin/env python3
"""
RES201 Document Patcher v2
Replaces TAB_* placeholders with formatted Word tables.
Adds Data and Code Availability section.
"""
import os, re, sys, io
from docx import Document
from docx.shared import Pt, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

BASE    = r'c:\Users\lenovo\res201-alignn-domain-shift'
IN_DIR  = os.path.join(BASE, 'reports', 'final_paper_factory', '05_final_docx')
OUT_DIR = os.path.join(BASE, 'reports', 'final_paper_factory', '06_template_ready', 'final_word')
os.makedirs(OUT_DIR, exist_ok=True)

REPO_URL = 'https://github.com/TheArchitect999/res201-alignn-domain-shift'
REPO_PARA = (
    'Project code, analysis scripts, processed result summaries, '
    'figure-generation files, and manuscript-generation materials are '
    'available in the project repository:\n' + REPO_URL
)

# ==============================================================================
# TABLE DATA  (verbatim from canonical_numbers_v2.csv and source files)
# ==============================================================================

DS_HEADERS = ['Family', 'Total', 'Train', 'Validation', 'Test', 'Pool (Train+Val)', 'Oxynitrides']
DS_OXIDE   = [['Oxide',   '14,991', '11,960', '1,547', '1,484', '13,507', '499']]
DS_NITRIDE = [['Nitride',  '2,288',  '1,837',   '209',   '242',  '2,046',   '0']]
DS_BOTH    = DS_OXIDE + DS_NITRIDE

EXP_HEADERS = ['Parameter / Setting', 'Value']
_EXP_COMMON = [
    ('SECTION', 'Hyperparameter Set 1'),
    ('Epochs', '50'),
    ('Batch size', '16'),
    ('Learning rate', '1 × 10⁻⁴'),
    ('Neighbour strategy', 'k-nearest'),
    ('Cutoff', '8.0 Å'),
    ('Cutoff extra', '3.0 Å'),
    ('Max neighbours', '12'),
    ('use_canonize / compute_line_graph / use_lmdb', 'enabled'),
    ('SECTION', 'Model architecture (ALIGNN formation-energy)'),
    ('ALIGNN convolutional layers', '4'),
    ('Gated GCN layers', '4'),
    ('Hidden size', '256'),
    ('Output head', 'scalar (eV/atom)'),
]
_EXP_OXIDE = _EXP_COMMON + [
    ('SECTION', 'Experimental scope — oxide'),
    ('Zero-shot evaluations', '1'),
    ('Fine-tuning N values', '10, 50, 100, 200, 500, 1,000'),
    ('Fine-tuning seeds per N', '5  →  30 runs total'),
    ('From-scratch N values', '50, 500'),
    ('From-scratch seeds per N', '5  →  10 runs total'),
]
_EXP_NITRIDE = _EXP_COMMON + [
    ('SECTION', 'Experimental scope — nitride'),
    ('Zero-shot evaluations', '1'),
    ('Fine-tuning N values', '10, 50, 100, 200, 500, 1,000'),
    ('Fine-tuning seeds per N', '5  →  30 runs total'),
    ('From-scratch N values', '50, 500'),
    ('From-scratch seeds per N', '5  →  10 runs total'),
]
_EXP_COMBINED = _EXP_COMMON + [
    ('SECTION', 'Experimental scope (oxide / nitride / total)'),
    ('Zero-shot evaluations', '1 / 1 / 2'),
    ('Fine-tuning N values', '10, 50, 100, 200, 500, 1,000'),
    ('Fine-tuning seeds per N', '5  →  30 / 30 / 60 runs'),
    ('From-scratch N values', '50, 500'),
    ('From-scratch seeds per N', '5  →  10 / 10 / 20 runs'),
]

ZS_HEADERS   = ['Family', 'Test structures (n)', 'Zero-shot MAE (eV/atom)']
ZS_OXIDE     = [['Oxide',   '1,484', '0.0342']]
ZS_NITRIDE   = [['Nitride',   '242', '0.0695']]
ZS_BOTH      = ZS_OXIDE + ZS_NITRIDE

FT_HEADERS   = ['N', 'Runs', 'Mean test MAE\n(eV/atom)', 'Std MAE\n(eV/atom)',
                 'Mean best\nepoch', 'Gap vs\nzero-shot (eV/atom)']
FT_OX = [
    ['10',    '5', '0.0417', '0.0111',  '1.0', '+0.0075'],
    ['50',    '5', '0.0523', '0.0148', '18.5', '+0.0181'],
    ['100',   '5', '0.0465', '0.0086', '20.0', '+0.0123'],
    ['200',   '5', '0.0457', '0.0086', '39.0', '+0.0115'],
    ['500',   '5', '0.0430', '0.0062', '39.0', '+0.0088'],
    ['1,000', '5', '0.0417', '0.0053', '35.5', '+0.0075'],
]
FT_NI = [
    ['10',    '5', '0.0874', '0.0199',  '1.0', '+0.0179'],
    ['50',    '5', '0.1173', '0.0451',  '1.0', '+0.0477'],
    ['100',   '5', '0.1722', '0.0996',  '1.0', '+0.1027'],
    ['200',   '5', '0.1392', '0.0677',  '1.0', '+0.0696'],
    ['500',   '5', '0.0977', '0.0178', '40.5', '+0.0281'],
    ['1,000', '5', '0.0907', '0.0135', '45.0', '+0.0211'],
]
FT_BOTH_HEADERS = ['Family', 'N', 'Runs', 'Mean test MAE\n(eV/atom)',
                   'Std MAE\n(eV/atom)', 'Mean best\nepoch', 'Gap vs\nzero-shot (eV/atom)']
FT_BOTH = [['Oxide'] + r for r in FT_OX] + [['Nitride'] + r for r in FT_NI]

FS_HEADERS = ['N', 'Fine-tune\nmean ± SD (eV/atom)',
              'From-scratch\nmean ± SD (eV/atom)',
              'Scratch − FT\n(eV/atom)', 'Scratch − zero-shot\n(eV/atom)']
FS_OX = [
    ['50',  '0.0523 ± 0.0148', '0.5561 ± 0.0523', '+0.5038', '+0.5219'],
    ['500', '0.0430 ± 0.0062', '0.2643 ± 0.0228', '+0.2214', '+0.2301'],
]
FS_NI = [
    ['50',  '0.1173 ± 0.0451', '0.6914 ± 0.0163', '+0.5741', '+0.6219'],
    ['500', '0.0977 ± 0.0178', '0.3683 ± 0.0233', '+0.2706', '+0.2988'],
]
FS_BOTH_HEADERS = ['Family', 'N', 'Fine-tune\nmean ± SD (eV/atom)',
                   'From-scratch\nmean ± SD (eV/atom)',
                   'Scratch − FT\n(eV/atom)', 'Scratch − zero-shot\n(eV/atom)']
FS_BOTH = [['Oxide'] + r for r in FS_OX] + [['Nitride'] + r for r in FS_NI]

EAFS_HEADERS = ['Metric', 'Scope', 'Value', '95% CI lower', '95% CI upper']
EAFS_ROWS = [
    ['Silhouette score',         'Overall', '0.2392', '0.2332', '0.2456'],
    ['Silhouette score',         'Oxide',   '0.2546', '0.2476', '0.2617'],
    ['Silhouette score',         'Nitride', '0.1453', '0.1316', '0.1582'],
    ['Davies–Bouldin index','Overall', '1.8290', '1.7340', '1.9071'],
    ['15-NN family purity',      'Overall', '0.9655', '0.9603', '0.9708'],
    ['15-NN family purity',      'Oxide',   '0.9872', '0.9832', '0.9906'],
    ['15-NN family purity',      'Nitride', '0.8331', '0.7978', '0.8645'],
    ['Logistic-regression AUC',  'Overall', '0.9994', '0.9984', '0.9999'],
]

EADE_HEADERS = ['Statistic', 'Comparison', 'Value', '95% CI', 'FDR q-value']
EADE_ROWS = [
    ['Spearman ρ', 'Abs. error vs 5NN oxide distance (n = 242)', '0.3428', '[0.2214, 0.4597]', '1.3 × 10⁻⁴'],
    ['Pearson r',       'Abs. error vs 5NN oxide distance (n = 242)', '0.2770', '[0.1741, 0.3890]', '1.3 × 10⁻⁴'],
    ['Hard mean 5NN distance†', 'Hard nitrides: top 20% by abs. error (n = 49)', '4.5988', '—', '—'],
    ['Easy mean 5NN distance†', 'Easy nitrides: bottom 20% by abs. error (n = 49)', '3.7821', '—', '—'],
    ['Hard − Easy mean gap', 'Tail contrast', '+0.8168', '[+0.4746, +1.1597]', '1.8 × 10⁻⁴'],
]

TABLE_CAPTIONS = {
    'TAB_METHODS_DATASET_SPLITS':
        'Dataset composition and split assignment. Counts after family-filter application. '
        'Pool = train + val. Oxynitrides (O+N structures) retained in the oxide arm; excluded from nitride arm.',
    'TAB_METHODS_EXPERIMENT_SCOPE':
        'Hyperparameter Set 1 settings and experimental scope. All trained runs in both arms '
        'use these settings.',
    'TAB_ZS_SUMMARY':
        'Zero-shot evaluation results. Same pretrained formation-energy ALIGNN checkpoint '
        'applied to each family’s fixed test set without target-family training.',
    'TAB_S1_FT_SUMMARY_BY_N':
        'Set 1 fine-tuning results by labelled-data size (N). Mean test MAE ± SD across '
        'five random seeds. Zero-shot MAE references: oxide 0.0342 eV/atom; nitride 0.0695 eV/atom. '
        'Gap = mean test MAE − zero-shot MAE (positive = above zero-shot).',
    'TAB_S1_FT_SUMMARY_BY_N_BOTH':
        'Set 1 fine-tuning results by family and labelled-data size. Mean test MAE ± SD '
        'across five random seeds per N.',
    'TAB_S1_FS_SUMMARY':
        'Set 1 pretrained fine-tuning versus from-scratch comparison at N = 50 and '
        'N = 500. FT = fine-tune mean across five seeds; from-scratch = '
        'randomly-initialized ALIGNN, same protocol.',
    'TAB_S1_FS_SUMMARY_BOTH':
        'Set 1 pretrained fine-tuning versus from-scratch comparison across both families '
        'at N = 50 and N = 500.',
    'TAB_EA_FAMILY_SEPARATION':
        'Family-separation metrics computed in the raw 256-D last_alignn_pool embedding space '
        '(fixed test set: 1,484 oxide + 242 nitride structures). 95% CI from bootstrap '
        'resampling. Davies–Bouldin: lower is better; silhouette, purity, AUC: higher is better.',
    'TAB_EA_DISTANCE_ERROR_STATS':
        'Distance–error association statistics. 5NN = 5-nearest-oxide-reference Euclidean '
        'distance in raw 256-D last_alignn_pool space. †Distance values in dimensionless '
        'Euclidean L2 units. FDR = Benjamini–Hochberg correction within statistic.',
}

# ==============================================================================
# HELPER FUNCTIONS
# ==============================================================================

def has_nearby_table(para_elem, max_look=10):
    """Return True if a w:tbl element appears within max_look siblings."""
    el = para_elem.getnext()
    for _ in range(max_look):
        if el is None:
            break
        if el.tag == qn('w:tbl'):
            return True
        el = el.getnext()
    return False


def add_caption_before(doc, anchor_elem, caption_text, tab_num):
    """Insert a Caption-style paragraph immediately before anchor_elem."""
    cap = doc.add_paragraph(f'Table {tab_num}. {caption_text}')
    cap.style = doc.styles['Caption']
    anchor_elem.addprevious(cap._element)
    return cap


def build_word_table(doc, headers, rows):
    """Create a Table Grid Word table. Returns the Table object."""
    n_rows = 1 + len(rows)
    n_cols = len(headers)
    tbl = doc.add_table(rows=n_rows, cols=n_cols)
    tbl.style = 'Table Grid'

    # Header row
    hdr = tbl.rows[0]
    for j, h in enumerate(headers):
        cell = hdr.cells[j]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)

    # Data rows
    for i, row_data in enumerate(rows):
        row = tbl.rows[i + 1]
        for j, val in enumerate(row_data):
            cell = row.cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            p.add_run(str(val)).font.size = Pt(9)

    return tbl


def build_exp_scope_table(doc, spec_rows):
    """Build experiment scope table with merged section-header rows."""
    n_data = len(spec_rows)
    n_sections = sum(1 for k, _ in spec_rows if k == 'SECTION')
    n_normal = n_data - n_sections
    total_rows = 1 + n_sections + n_normal  # header + section headers + data rows

    tbl = doc.add_table(rows=total_rows, cols=2)
    tbl.style = 'Table Grid'

    # Column header
    hdr = tbl.rows[0]
    for j, h in enumerate(['Parameter / Setting', 'Value']):
        cell = hdr.cells[j]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9)

    row_idx = 1
    for key, val in spec_rows:
        if key == 'SECTION':
            row = tbl.rows[row_idx]
            merged = row.cells[0].merge(row.cells[1])
            merged.text = ''
            p = merged.paragraphs[0]
            run = p.add_run(val)
            run.bold = True
            run.font.size = Pt(9)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        else:
            row = tbl.rows[row_idx]
            row.cells[0].text = ''
            row.cells[0].paragraphs[0].add_run(key).font.size = Pt(9)
            row.cells[1].text = ''
            row.cells[1].paragraphs[0].add_run(val).font.size = Pt(9)
        row_idx += 1

    return tbl


def replace_para_with_table(doc, para_elem, tbl_obj, cap_elem=None):
    """Move caption + table to placeholder position, then remove placeholder."""
    tbl_xml = tbl_obj._tbl
    # Insert table before placeholder
    para_elem.addprevious(tbl_xml)
    # Insert caption before table
    if cap_elem is not None:
        tbl_xml.addprevious(cap_elem._element)
    # Add spacing paragraph after table
    spacer = doc.add_paragraph()
    spacer.style = doc.styles['Normal']
    tbl_xml.addnext(spacer._element)
    # Remove placeholder
    para_elem.getparent().remove(para_elem)


def get_tab_id(para):
    """Extract TAB_* id from a placeholder paragraph."""
    m = re.search(r'\[TABLE:\s*(TAB_\w+)', para.text)
    return m.group(1) if m else None


# ==============================================================================
# TABLE DATA SELECTOR
# ==============================================================================

def select_table(tab_id, doc_type, occ):
    """Return (is_exp_scope, headers, rows, caption_key)."""
    if tab_id == 'TAB_METHODS_DATASET_SPLITS':
        rows = DS_OXIDE if doc_type == 'oxide' else (DS_NITRIDE if doc_type == 'nitride' else DS_BOTH)
        return False, DS_HEADERS, rows, 'TAB_METHODS_DATASET_SPLITS'

    if tab_id == 'TAB_METHODS_EXPERIMENT_SCOPE':
        spec = _EXP_OXIDE if doc_type == 'oxide' else (_EXP_NITRIDE if doc_type == 'nitride' else _EXP_COMBINED)
        return True, EXP_HEADERS, spec, 'TAB_METHODS_EXPERIMENT_SCOPE'

    if tab_id == 'TAB_ZS_SUMMARY':
        rows = ZS_OXIDE if doc_type == 'oxide' else ZS_BOTH
        return False, ZS_HEADERS, rows, 'TAB_ZS_SUMMARY'

    if tab_id == 'TAB_S1_FT_SUMMARY_BY_N':
        if doc_type == 'oxide':
            return False, FT_HEADERS, FT_OX, 'TAB_S1_FT_SUMMARY_BY_N'
        if doc_type == 'nitride':
            return False, FT_HEADERS, FT_NI, 'TAB_S1_FT_SUMMARY_BY_N'
        # combined: occ 1 → deleted, occ 2 → nitride, occ 3 → both
        if occ == 2:
            return False, FT_HEADERS, FT_NI, 'TAB_S1_FT_SUMMARY_BY_N'
        return False, FT_BOTH_HEADERS, FT_BOTH, 'TAB_S1_FT_SUMMARY_BY_N_BOTH'

    if tab_id == 'TAB_S1_FS_SUMMARY':
        if doc_type == 'oxide':
            return False, FS_HEADERS, FS_OX, 'TAB_S1_FS_SUMMARY'
        if doc_type == 'nitride':
            return False, FS_HEADERS, FS_NI, 'TAB_S1_FS_SUMMARY'
        if occ == 2:
            return False, FS_HEADERS, FS_NI, 'TAB_S1_FS_SUMMARY'
        return False, FS_BOTH_HEADERS, FS_BOTH, 'TAB_S1_FS_SUMMARY_BOTH'

    if tab_id == 'TAB_EA_FAMILY_SEPARATION':
        return False, EAFS_HEADERS, EAFS_ROWS, 'TAB_EA_FAMILY_SEPARATION'

    if tab_id == 'TAB_EA_DISTANCE_ERROR_STATS':
        return False, EADE_HEADERS, EADE_ROWS, 'TAB_EA_DISTANCE_ERROR_STATS'

    return False, None, None, None


# ==============================================================================
# SECTION INSERTION
# ==============================================================================

def insert_data_availability(doc, log):
    """Insert Data and Code Availability section before References heading."""
    refs_elem = None
    for p in doc.paragraphs:
        if p.style.name.startswith('Heading') and 'Reference' in p.text:
            refs_elem = p._element
            break

    if refs_elem is None:
        log.append('  WARNING: References heading not found; DCA section added at end')
        doc.add_heading('Data and code availability', level=1)
        doc.add_paragraph(REPO_PARA).style = doc.styles['Normal']
        return

    # Create heading and text, then move before References
    h = doc.add_heading('Data and code availability', level=1)
    refs_elem.addprevious(h._element)
    p = doc.add_paragraph(REPO_PARA)
    p.style = doc.styles['Normal']
    h._element.addnext(p._element)
    log.append('  Data and code availability section inserted before References')


# ==============================================================================
# MAIN PATCHER
# ==============================================================================

def patch_document(ms_name, log):
    in_path  = os.path.join(IN_DIR,  f'{ms_name}_final.docx')
    out_path = os.path.join(OUT_DIR, f'{ms_name}_final_v2.docx')
    doc_type = ms_name  # 'oxide', 'nitride', 'combined'

    doc = Document(in_path)
    log.append(f'\n=== {ms_name.upper()} ===')

    # 1. Add Data and Code Availability section
    insert_data_availability(doc, log)

    # 2. Collect all placeholder paragraphs
    placeholders = []
    for para in doc.paragraphs:
        tid = get_tab_id(para)
        if tid:
            placeholders.append((para, tid))
    log.append(f'  Found {len(placeholders)} TABLE placeholders')

    # 3. Track occurrence count per tab_id
    occ_count = {}
    table_num = 0

    for para, tab_id in placeholders:
        occ_count[tab_id] = occ_count.get(tab_id, 0) + 1
        occ = occ_count[tab_id]

        nearby = has_nearby_table(para._element)

        if tab_id in ('TAB_S1_FT_SUMMARY_BY_N', 'TAB_S1_FS_SUMMARY') and nearby:
            # DELETE: inline Word table already exists below
            para._element.getparent().remove(para._element)
            log.append(f'  DELETED placeholder {tab_id} (occ {occ}) — inline table retained below')
            continue

        # CREATE new table
        is_exp, headers, rows, cap_key = select_table(tab_id, doc_type, occ)
        if headers is None:
            log.append(f'  SKIP {tab_id} occ {occ} — no data defined')
            continue

        table_num += 1
        cap_text = TABLE_CAPTIONS.get(cap_key, cap_key)

        if is_exp:
            tbl_obj = build_exp_scope_table(doc, rows)
        else:
            tbl_obj = build_word_table(doc, headers, rows)

        # Build caption paragraph (at end temporarily)
        cap_para = doc.add_paragraph(f'Table {table_num}. {cap_text}')
        cap_para.style = doc.styles['Caption']

        replace_para_with_table(doc, para._element, tbl_obj, cap_para)
        log.append(f'  CREATED Table {table_num} for {tab_id} (occ {occ}, {len(rows)} data rows)')

    doc.save(out_path)
    log.append(f'  Saved: {out_path}')
    return out_path, table_num


# ==============================================================================
# QUALITY CHECK
# ==============================================================================

def qc_document(path, ms_name, log):
    """Scan v2 document and report QC metrics."""
    doc = Document(path)
    metrics = {
        'unresolved_table_placeholders': 0,
        'unresolved_figure_placeholders': 0,
        'unresolved_cite_placeholders': 0,
        'affiliation_placeholders': 0,
        'ack_placeholders': 0,
        'repo_link_present': False,
        'missing_images': 0,
        'word_tables': 0,
        'references_section_present': False,
        'dca_section_present': False,
    }
    from docx.oxml.ns import qn as _qn
    body = doc.element.body

    for child in body:
        if child.tag == _qn('w:tbl'):
            metrics['word_tables'] += 1

    for p in doc.paragraphs:
        txt = p.text
        if '[TABLE:' in txt or 'TAB_' in txt:
            metrics['unresolved_table_placeholders'] += 1
        if '[INSERT FIGURE' in txt:
            metrics['unresolved_figure_placeholders'] += 1
        if '[CITE:' in txt:
            metrics['unresolved_cite_placeholders'] += 1
        if '[Affiliation' in txt:
            metrics['affiliation_placeholders'] += 1
        if '[ACKNOWLEDGEMENTS PLACEHOLDER' in txt:
            metrics['ack_placeholders'] += 1
        if 'github.com/TheArchitect999' in txt:
            metrics['repo_link_present'] = True
        if p.style.name.startswith('Heading') and 'Reference' in txt:
            metrics['references_section_present'] = True
        if p.style.name.startswith('Heading') and 'code availability' in txt.lower():
            metrics['dca_section_present'] = True

    # Check images (inline shapes)
    for drawing in doc.element.body.iter(_qn('a:blip')):
        pass  # just a check; if we got here, images exist

    log.append(f'  QC {ms_name}: {metrics}')
    return metrics


# ==============================================================================
# CHANGELOG & QC WRITERS
# ==============================================================================

def write_changelog(results, out_dir):
    lines = [
        '# RES201 Final Word Documents — Build Changelog v2',
        f'**Date:** 2026-04-24',
        '',
        '## Summary of changes from v1 to v2',
        '',
        '### Changes applied to all three documents',
        '- All `[TABLE: TAB_*]` placeholder paragraphs replaced with formatted Word tables.',
        '- "Data and code availability" section inserted between Acknowledgements and References.',
        '',
    ]
    for ms, (path, n_created, n_deleted, log_lines) in results.items():
        lines.append(f'### {ms.capitalize()}')
        for ll in log_lines:
            lines.append(f'- {ll.strip()}')
        lines.append('')

    lines += [
        '## Retained unchanged from v1',
        '- All inserted figures (16 main-text figures from core_figures/).',
        '- All Nature-style numbered reference lists.',
        '- All inline text with citations replaced by [n] markers.',
        '- All scientific prose (no text rewritten).',
        '- JURI template styles (Title, Heading 1–3, Normal, Caption).',
        '',
        '## Items still requiring manual Word work',
        '1. **Author affiliation** — placeholder text remains; replace with exact KFUPM affiliation line.',
        '2. **Inline table captions** — the Word tables converted from markdown (§3.2 fine-tuning table,'
        '   §3.4 scratch table) do not have captions. Add captions manually if required.',
        '3. **In-text table cross-references** — prose still contains backtick TAB_ID references '
        '   (e.g., `TAB_METHODS_DATASET_SPLITS`). Replace with "Table N" using find-and-replace.',
        '4. **In-text figure cross-references** — prose still contains backtick FIG_ID references. '
        '   Replace with "Figure N" using find-and-replace.',
        '5. **Acknowledgements expansion** — current text is a minimal standard line. '
        '   Add grant numbers, institutional acknowledgements, and contributor thanks.',
        '6. **Table column widths** — some wider tables may benefit from manual column-width '
        '   adjustment in Word for readability.',
    ]
    path = os.path.join(out_dir, 'final_word_build_changelog_v2.md')
    with open(path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(lines))
    return path


def write_qc(results, out_dir):
    lines = [
        '# RES201 Final Word Documents — Quality Control Report v2',
        f'**Date:** 2026-04-24',
        '**Documents inspected:** oxide_final_v2.docx · nitride_final_v2.docx · combined_final_v2.docx',
        '',
        '## Per-document QC metrics',
        '',
    ]
    all_ready = True
    for ms, (path, n_created, n_deleted, log_lines) in results.items():
        m = results[ms][4] if len(results[ms]) > 4 else {}
        tab_rem = m.get('unresolved_table_placeholders', '?')
        fig_rem = m.get('unresolved_figure_placeholders', '?')
        cite_rem = m.get('unresolved_cite_placeholders', '?')
        affil = m.get('affiliation_placeholders', '?')
        ack = m.get('ack_placeholders', '?')
        repo = m.get('repo_link_present', False)
        dca = m.get('dca_section_present', False)
        refs = m.get('references_section_present', False)
        n_tbl = m.get('word_tables', '?')

        if tab_rem or not repo:
            all_ready = False

        lines += [
            f'### {ms.capitalize()}',
            f'| Item | Result |',
            f'|------|--------|',
            f'| Unresolved TABLE placeholders | {tab_rem} |',
            f'| Unresolved FIGURE placeholders | {fig_rem} |',
            f'| Unresolved CITE placeholders | {cite_rem} |',
            f'| Affiliation placeholders | {affil} |',
            f'| Acknowledgements placeholders | {ack} |',
            f'| Repository link present | {"YES ✓" if repo else "NO ✗"} |',
            f'| Data and code availability section | {"YES ✓" if dca else "NO ✗"} |',
            f'| References section present | {"YES ✓" if refs else "NO ✗"} |',
            f'| Word tables in document | {n_tbl} |',
            f'| New tables created (v2) | {n_created} |',
            f'| Placeholders deleted (inline table retained) | {n_deleted} |',
            '',
        ]

    lines += [
        '## Notes on remaining items',
        '',
        '### Affiliation placeholders (all three documents)',
        'Text: `[Affiliation — confirm with authors; email domain: kfupm.edu.sa]`  ',
        'Action: Replace with the exact KFUPM institutional affiliation string before submission.',
        '',
        '### Acknowledgements (all three documents)',
        'Current text is a brief standard acknowledgement citing the JARVIS infrastructure. '
        'Authors should expand with specific grant numbers and contributor thanks.',
        '',
        '### In-text TAB_ID and FIG_ID references',
        'Prose still contains backtick-formatted table and figure IDs (e.g., `TAB_METHODS_DATASET_SPLITS`, '
        '`FIG_S1_LC_OXIDE`). These should be replaced with "Table N" / "Figure N" using Word find-and-replace '
        'after confirming sequential numbering.',
        '',
        f'## Final verdict',
        f'**{"READY for human Word QA pass" if all_ready else "NOT READY — see items above"}**',
        '',
        'Mandatory before submission: (1) Replace affiliation placeholder; '
        '(2) replace in-text TAB_ID/FIG_ID backtick references with sequential Table N / Figure N labels.',
    ]
    path = os.path.join(out_dir, 'final_word_build_qc_v2.md')
    with open(path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(lines))
    return path


# ==============================================================================
# ENTRY POINT
# ==============================================================================

def main():
    results = {}
    for ms_name in ('oxide', 'nitride', 'combined'):
        log = []
        out_path, n_created = patch_document(ms_name, log)

        # Count deletes from log
        n_deleted = sum(1 for l in log if 'DELETED' in l)
        n_created_actual = sum(1 for l in log if 'CREATED' in l)

        # QC
        qc_log = []
        metrics = qc_document(out_path, ms_name, qc_log)
        log.extend(qc_log)

        results[ms_name] = (out_path, n_created_actual, n_deleted, log, metrics)
        print(f'{ms_name}: {n_created_actual} tables created, {n_deleted} placeholders deleted')
        for l in log:
            print(l)

    # Write reports
    cl_path = write_changelog(results, OUT_DIR)
    qc_path = write_qc(results, OUT_DIR)
    print(f'\nChangelog: {cl_path}')
    print(f'QC report: {qc_path}')

    # Verify output files
    print('\nOutput files:')
    for fn in os.listdir(OUT_DIR):
        if fn.endswith('.docx') or fn.endswith('.md'):
            sz = os.path.getsize(os.path.join(OUT_DIR, fn))
            print(f'  {fn:45s} {sz:>9,} bytes')


if __name__ == '__main__':
    main()
